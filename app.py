import streamlit as st
from content_agent import ContentAgent
from docx import Document

st.set_page_config(page_title="Smart Research & Summary Agent", layout="centered")
st.title("📝 Smart Research & Summary Agent")

# 用户输入 API Key
api_key = st.text_input("请输入你的 OpenAI API Key:", type="password")

# 用户输入搜索主题
query = st.text_input("请输入搜索关键词或主题:")

if st.button("生成摘要"):
    if not api_key or not query:
        st.warning("请填写 API Key 和搜索关键词！")
    else:
        agent = ContentAgent(api_key)
        with st.spinner("正在抓取资料并生成摘要..."):
            result = agent.generate_summary_from_query(query)

        # 显示摘要
        st.markdown("### 摘要结果")
        st.write(result)

        # 提供下载 Word 文件
        doc = Document()
        doc.add_heading(query, 0)
        doc.add_paragraph(result)
        file_name = f"{query}_summary.docx"
        doc.save(file_name)
        st.success(f"摘要已生成并保存为 {file_name}")
